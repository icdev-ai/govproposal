#!/usr/bin/env python3
# CUI // SP-PROPIN
# Controlled by: GovProposal Portal
# CUI Category: PROPIN
# Distribution: D
# POC: GovProposal System Administrator
"""Submission Packager — Final document assembly and packaging for GovProposal.

Assembles final proposal packages including:
  - Section status validation (all must be 'final')
  - White team QC checks
  - CAG pre-export classification guard check
  - Table of contents generation
  - Acronym list extraction and expansion
  - Cross-reference validation
  - Classification marking application
  - Per-volume file packaging with naming convention
  - Submission checklist generation

Usage:
    python tools/production/submission_packager.py --package --proposal-id PROP-001 --output /tmp/out --json
    python tools/production/submission_packager.py --validate --proposal-id PROP-001 --json
    python tools/production/submission_packager.py --acronyms --proposal-id PROP-001 --json
    python tools/production/submission_packager.py --cross-refs --proposal-id PROP-001 --json
    python tools/production/submission_packager.py --status --proposal-id PROP-001 --json
"""

import json
import os
import re
import secrets
import sqlite3
import sys
from datetime import datetime, timezone
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent.parent
DB_PATH = Path(os.environ.get(
    "GOVPROPOSAL_DB_PATH", str(BASE_DIR / "data" / "govproposal.db")
))

# Optional YAML import
try:
    import yaml  # noqa: F401
except ImportError:
    yaml = None

# Optional python-docx import for Word export
try:
    import docx as python_docx  # type: ignore
    from docx.shared import Inches, Pt, Cm, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
    from docx.enum.section import WD_ORIENT  # type: ignore
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _pkg_id():
    """Generate a package ID: PKG- followed by 12 hex characters."""
    return "PKG-" + secrets.token_hex(6)


def _acr_id():
    """Generate an acronym ID: ACR- followed by 12 hex characters."""
    return "ACR-" + secrets.token_hex(6)


def _now():
    """Return current UTC timestamp as ISO-8601 string."""
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _get_db(db_path=None):
    """Open a database connection with WAL mode and foreign keys enabled.

    Args:
        db_path: Optional path override.  Falls back to DB_PATH.

    Returns:
        sqlite3.Connection with row_factory set to sqlite3.Row.
    """
    path = str(db_path or DB_PATH)
    conn = sqlite3.connect(path)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def _audit(conn, event_type, action, entity_type=None, entity_id=None,
           details=None):
    """Write an append-only audit trail record.

    Args:
        conn: Active database connection.
        event_type: Category of event.
        action: Human-readable description.
        entity_type: Type of entity affected.
        entity_id: ID of the affected entity.
        details: Optional JSON-serializable details dict.
    """
    conn.execute(
        "INSERT INTO audit_trail (event_type, actor, action, entity_type, "
        "entity_id, details, created_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            event_type,
            "submission_packager",
            action,
            entity_type,
            entity_id,
            json.dumps(details) if details else None,
            _now(),
        ),
    )


def _row_to_dict(row):
    """Convert a sqlite3.Row to a plain dict."""
    if row is None:
        return None
    return dict(row)


def _parse_json_field(value):
    """Safely parse a JSON string field."""
    if value is None:
        return None
    try:
        return json.loads(value)
    except (json.JSONDecodeError, TypeError):
        return value


def _load_proposal(conn, proposal_id):
    """Load proposal record and raise if not found."""
    row = conn.execute(
        "SELECT * FROM proposals WHERE id = ?", (proposal_id,)
    ).fetchone()
    if row is None:
        raise ValueError(f"Proposal not found: {proposal_id}")
    return _row_to_dict(row)


def _load_sections(conn, proposal_id):
    """Load all proposal sections ordered by volume and section number."""
    rows = conn.execute(
        "SELECT * FROM proposal_sections WHERE proposal_id = ? "
        "ORDER BY volume, section_number",
        (proposal_id,),
    ).fetchall()
    return [_row_to_dict(r) for r in rows]


def _safe_filename(text, max_len=50):
    """Sanitize a string for use as a filename component.

    Args:
        text: Raw string.
        max_len: Maximum length.

    Returns:
        Sanitized string with only alphanumerics, hyphens, underscores.
    """
    sanitized = re.sub(r"[^\w\-]", "_", text.strip())
    sanitized = re.sub(r"_+", "_", sanitized).strip("_")
    return sanitized[:max_len]


# ---------------------------------------------------------------------------
# Core Functions
# ---------------------------------------------------------------------------

def package_proposal(proposal_id, output_dir, db_path=None):
    """Assemble the final proposal package for submission.

    Steps:
      1. Validate all sections are 'final' status
      2. Run White team QC checks
      3. Run CAG pre-export check
      4. Generate table of contents
      5. Build acronym list from all sections
      6. Validate cross-references
      7. Apply classification markings to every page
      8. Package into per-volume files with naming convention
      9. Generate submission checklist

    Args:
        proposal_id: Proposal ID to package.
        output_dir: Directory to write package files into.
        db_path: Optional database path override.

    Returns:
        dict with package manifest: file list, sizes, compliance status,
        checklist items.

    Raises:
        ValueError: If proposal not found or pre-conditions fail.
    """
    conn = _get_db(db_path)
    try:
        proposal = _load_proposal(conn, proposal_id)
        sections = _load_sections(conn, proposal_id)
        if not sections:
            raise ValueError(
                f"No sections found for proposal {proposal_id}"
            )

        classification = proposal.get("classification") or "CUI // SP-PROPIN"
        opp_id = proposal.get("opportunity_id") or "UNKNOWN"
        sol_num = ""
        if opp_id != "UNKNOWN":
            opp_row = conn.execute(
                "SELECT solicitation_number, title FROM opportunities "
                "WHERE id = ?", (opp_id,)
            ).fetchone()
            if opp_row:
                sol_num = opp_row["solicitation_number"] or ""

        issues = []
        warnings = []
        files_written = []

        # --- Step 1: Validate all sections are 'final' ---
        non_final = [s for s in sections if s.get("status") != "final"]
        if non_final:
            non_final_ids = [
                f"{s['section_number']} ({s['status']})"
                for s in non_final
            ]
            issues.append(
                f"Non-final sections: {', '.join(non_final_ids)}"
            )

        # --- Step 2: White team QC check ---
        latest_white = conn.execute(
            "SELECT * FROM proposal_reviews "
            "WHERE proposal_id = ? AND review_type = 'white' "
            "ORDER BY reviewed_at DESC LIMIT 1",
            (proposal_id,),
        ).fetchone()
        if latest_white:
            white = _row_to_dict(latest_white)
            white_score = white.get("overall_score") or 0.0
            if white_score < 0.95:
                warnings.append(
                    f"White team score {white_score:.1%} is below 95% "
                    f"threshold. Review deficiencies before submission."
                )
            white_defs = _parse_json_field(white.get("deficiencies")) or []
            if white_defs:
                issues.append(
                    f"White team deficiencies unresolved: {len(white_defs)}"
                )
        else:
            warnings.append("No White team review found. QC not verified.")

        # --- Step 3: CAG pre-export check ---
        cag_status = proposal.get("cag_status") or "pending"
        open_alerts = conn.execute(
            "SELECT COUNT(*) as cnt FROM cag_alerts "
            "WHERE proposal_id = ? AND status IN ('open', 'quarantined')",
            (proposal_id,),
        ).fetchone()
        open_alert_count = open_alerts["cnt"] if open_alerts else 0
        if cag_status == "blocked" or cag_status == "quarantined":
            issues.append(
                f"CAG status is '{cag_status}'. Cannot export."
            )
        if open_alert_count > 0:
            issues.append(
                f"{open_alert_count} open CAG alert(s). Resolve before export."
            )

        # --- Step 4: Generate table of contents ---
        toc_entries = []
        for s in sections:
            toc_entries.append({
                "volume": s["volume"],
                "section_number": s["section_number"],
                "section_title": s["section_title"],
                "page_count": s.get("page_count") or 0,
            })

        # --- Step 5: Build acronym list ---
        acronyms_found = generate_acronym_list(
            proposal_id, db_path=db_path
        )

        # --- Step 6: Validate cross-references ---
        xref_results = check_cross_references(
            proposal_id, db_path=db_path
        )
        broken_refs = xref_results.get("broken_references", [])
        if broken_refs:
            warnings.append(
                f"{len(broken_refs)} potentially broken cross-reference(s)."
            )

        # --- Step 7 & 8: Write per-volume files ---
        out_path = Path(output_dir)
        out_path.mkdir(parents=True, exist_ok=True)

        volumes_defined = _parse_json_field(proposal.get("volumes")) or [
            "technical", "management", "past_performance", "cost"
        ]
        volume_sections = {}
        for s in sections:
            vol = s.get("volume", "technical")
            volume_sections.setdefault(vol, []).append(s)

        sol_slug = _safe_filename(sol_num) if sol_num else "proposal"

        for vol_name in volumes_defined:
            vol_secs = volume_sections.get(vol_name, [])
            if not vol_secs:
                continue

            filename = f"{sol_slug}_vol_{vol_name}.txt"
            filepath = out_path / filename

            lines = []
            lines.append(classification)
            lines.append("")
            lines.append(f"PROPOSAL: {proposal.get('title', proposal_id)}")
            lines.append(f"VOLUME: {vol_name.upper()}")
            if sol_num:
                lines.append(f"SOLICITATION: {sol_num}")
            lines.append(f"CLASSIFICATION: {classification}")
            lines.append("=" * 72)
            lines.append("")

            # Table of contents for this volume
            lines.append("TABLE OF CONTENTS")
            lines.append("-" * 40)
            for s in vol_secs:
                lines.append(
                    f"  {s['section_number']}  {s['section_title']}"
                )
            lines.append("")
            lines.append("=" * 72)

            # Section content
            for s in vol_secs:
                lines.append("")
                lines.append(classification)
                lines.append(
                    f"SECTION {s['section_number']}: "
                    f"{s['section_title'].upper()}"
                )
                lines.append("-" * 60)
                content = s.get("content") or "[CONTENT MISSING]"
                lines.append(content)
                lines.append("")
                lines.append(classification)
                lines.append("")

            # Footer
            lines.append("=" * 72)
            lines.append(classification)

            filepath.write_text("\n".join(lines), encoding="utf-8")
            file_size = filepath.stat().st_size
            files_written.append({
                "filename": filename,
                "volume": vol_name,
                "sections": len(vol_secs),
                "size_bytes": file_size,
                "path": str(filepath),
            })

        # Write acronym list file
        if acronyms_found.get("acronyms"):
            acr_filename = f"{sol_slug}_acronyms.txt"
            acr_path = out_path / acr_filename
            acr_lines = [classification, "", "ACRONYM LIST", "-" * 40]
            for a in acronyms_found["acronyms"]:
                acr_lines.append(
                    f"  {a['acronym']:12s} {a['expansion']}"
                )
            acr_lines.append("", classification)
            acr_path.write_text("\n".join(acr_lines), encoding="utf-8")
            files_written.append({
                "filename": acr_filename,
                "volume": "attachments",
                "sections": 0,
                "size_bytes": acr_path.stat().st_size,
                "path": str(acr_path),
            })

        # --- Step 9: Generate submission checklist ---
        checklist = _generate_checklist(
            proposal, sections, issues, warnings,
            files_written, latest_white is not None,
            open_alert_count, broken_refs,
        )

        # Write checklist file
        chk_filename = f"{sol_slug}_checklist.txt"
        chk_path = out_path / chk_filename
        chk_lines = [classification, "", "SUBMISSION CHECKLIST", "=" * 40]
        for item in checklist:
            status_mark = "[X]" if item["passed"] else "[ ]"
            chk_lines.append(f"  {status_mark} {item['item']}")
            if not item["passed"] and item.get("note"):
                chk_lines.append(f"       NOTE: {item['note']}")
        chk_lines.append("", classification)
        chk_path.write_text("\n".join(chk_lines), encoding="utf-8")
        files_written.append({
            "filename": chk_filename,
            "volume": "admin",
            "sections": 0,
            "size_bytes": chk_path.stat().st_size,
            "path": str(chk_path),
        })

        # Overall compliance status
        all_passed = all(item["passed"] for item in checklist)
        compliance_status = "READY" if all_passed else "NOT_READY"
        if issues:
            compliance_status = "BLOCKED"

        now = _now()
        result = {
            "proposal_id": proposal_id,
            "package_id": _pkg_id(),
            "compliance_status": compliance_status,
            "output_dir": str(out_path),
            "files": files_written,
            "total_files": len(files_written),
            "total_size_bytes": sum(f["size_bytes"] for f in files_written),
            "issues": issues,
            "warnings": warnings,
            "checklist": checklist,
            "toc": toc_entries,
            "acronym_count": len(acronyms_found.get("acronyms", [])),
            "cross_ref_issues": len(broken_refs),
            "packaged_at": now,
        }

        _audit(conn, "production.package",
               f"Packaged proposal {proposal_id}: {compliance_status}",
               "proposal", proposal_id,
               {"compliance_status": compliance_status,
                "files": len(files_written),
                "issues": len(issues)})
        conn.commit()
        return result
    finally:
        conn.close()


def _generate_checklist(proposal, sections, issues, warnings,
                        files, has_white_review, open_alerts, broken_refs):
    """Generate a submission pre-flight checklist.

    Args:
        proposal: Proposal dict.
        sections: List of section dicts.
        issues: List of blocking issue strings.
        warnings: List of warning strings.
        files: List of file dicts written.
        has_white_review: Whether White team review exists.
        open_alerts: Number of open CAG alerts.
        broken_refs: List of broken cross-references.

    Returns:
        list of checklist item dicts with 'item', 'passed', 'note'.
    """
    checklist = []

    # 1. All sections finalized
    non_final = [s for s in sections if s.get("status") != "final"]
    checklist.append({
        "item": "All sections in 'final' status",
        "passed": len(non_final) == 0,
        "note": (f"{len(non_final)} non-final section(s)"
                 if non_final else None),
    })

    # 2. White team review completed
    checklist.append({
        "item": "White team review completed",
        "passed": has_white_review,
        "note": None if has_white_review else "No White team review on file",
    })

    # 3. CAG clearance
    checklist.append({
        "item": "CAG clearance — no open alerts",
        "passed": open_alerts == 0,
        "note": (f"{open_alerts} open alert(s)"
                 if open_alerts > 0 else None),
    })

    # 4. Classification markings applied
    classification = proposal.get("classification") or ""
    has_marking = bool(classification)
    checklist.append({
        "item": "Classification markings defined",
        "passed": has_marking,
        "note": None if has_marking else "No classification marking set",
    })

    # 5. Cross-references validated
    checklist.append({
        "item": "Cross-references validated",
        "passed": len(broken_refs) == 0,
        "note": (f"{len(broken_refs)} potential issue(s)"
                 if broken_refs else None),
    })

    # 6. Volumes packaged
    checklist.append({
        "item": "Volume files generated",
        "passed": len(files) > 0,
        "note": f"{len(files)} file(s) written",
    })

    # 7. Page limits respected
    over_limit = [
        s for s in sections
        if s.get("page_limit") and s.get("page_count")
        and s["page_count"] > s["page_limit"]
    ]
    checklist.append({
        "item": "All sections within page limits",
        "passed": len(over_limit) == 0,
        "note": (f"{len(over_limit)} section(s) over limit"
                 if over_limit else None),
    })

    # 8. Proposal manager assigned
    has_pm = bool(proposal.get("assigned_pm"))
    checklist.append({
        "item": "Proposal Manager assigned",
        "passed": has_pm,
        "note": None if has_pm else "No PM assigned",
    })

    # 9. Due date set
    has_due = bool(proposal.get("due_date"))
    checklist.append({
        "item": "Submission due date set",
        "passed": has_due,
        "note": proposal.get("due_date") if has_due else "No due date",
    })

    # 10. No blocking issues
    checklist.append({
        "item": "No blocking issues",
        "passed": len(issues) == 0,
        "note": (f"{len(issues)} blocking issue(s)"
                 if issues else None),
    })

    return checklist


def validate_submission(proposal_id, db_path=None):
    """Run pre-submission validation without generating package files.

    Performs the same checks as package_proposal but does not write files.

    Args:
        proposal_id: Proposal ID to validate.
        db_path: Optional database path override.

    Returns:
        dict with validation results and checklist.

    Raises:
        ValueError: If proposal not found.
    """
    conn = _get_db(db_path)
    try:
        proposal = _load_proposal(conn, proposal_id)
        sections = _load_sections(conn, proposal_id)
        if not sections:
            raise ValueError(
                f"No sections found for proposal {proposal_id}"
            )

        issues = []
        warnings = []

        # Section status check
        non_final = [s for s in sections if s.get("status") != "final"]
        if non_final:
            issues.append(
                f"{len(non_final)} section(s) not in 'final' status"
            )

        # White team review
        latest_white = conn.execute(
            "SELECT overall_score FROM proposal_reviews "
            "WHERE proposal_id = ? AND review_type = 'white' "
            "ORDER BY reviewed_at DESC LIMIT 1",
            (proposal_id,),
        ).fetchone()
        has_white = latest_white is not None
        if not has_white:
            warnings.append("White team review not completed")

        # CAG check
        open_alerts_row = conn.execute(
            "SELECT COUNT(*) as cnt FROM cag_alerts "
            "WHERE proposal_id = ? AND status IN ('open', 'quarantined')",
            (proposal_id,),
        ).fetchone()
        open_alerts = open_alerts_row["cnt"] if open_alerts_row else 0
        if open_alerts > 0:
            issues.append(f"{open_alerts} open CAG alert(s)")

        # Cross-ref check
        xrefs = check_cross_references(proposal_id, db_path=db_path)
        broken = xrefs.get("broken_references", [])
        if broken:
            warnings.append(
                f"{len(broken)} potentially broken cross-reference(s)"
            )

        # Page limits
        over_limit = [
            s for s in sections
            if s.get("page_limit") and s.get("page_count")
            and s["page_count"] > s["page_limit"]
        ]
        if over_limit:
            issues.append(f"{len(over_limit)} section(s) over page limit")

        checklist = _generate_checklist(
            proposal, sections, issues, warnings,
            [], has_white, open_alerts, broken,
        )

        all_passed = all(item["passed"] for item in checklist)
        status = "READY" if all_passed else "NOT_READY"
        if issues:
            status = "BLOCKED"

        result = {
            "proposal_id": proposal_id,
            "validation_status": status,
            "section_count": len(sections),
            "non_final_sections": len(non_final),
            "open_cag_alerts": open_alerts,
            "cross_ref_issues": len(broken),
            "page_limit_violations": len(over_limit),
            "issues": issues,
            "warnings": warnings,
            "checklist": checklist,
        }

        _audit(conn, "production.validate",
               f"Validated proposal {proposal_id}: {status}",
               "proposal", proposal_id,
               {"status": status, "issues": len(issues)})
        conn.commit()
        return result
    finally:
        conn.close()


def generate_acronym_list(proposal_id, db_path=None):
    """Extract all acronyms from proposal sections, expand, and store.

    Scans all section content for uppercase 2-6 letter sequences,
    checks the acronyms table for known expansions, and stores any
    new discoveries.

    Args:
        proposal_id: Proposal ID to scan.
        db_path: Optional database path override.

    Returns:
        dict with acronyms list, new_count, existing_count.

    Raises:
        ValueError: If proposal not found.
    """
    conn = _get_db(db_path)
    try:
        proposal = _load_proposal(conn, proposal_id)
        sections = _load_sections(conn, proposal_id)

        # Common English words that look like acronyms
        non_acronyms = {
            "THE", "AND", "FOR", "NOT", "BUT", "ALL", "HAS", "WAS",
            "ARE", "CAN", "OUR", "HIS", "HER", "ITS", "MAY", "USE",
            "SET", "NEW", "OLD", "ONE", "TWO", "PER", "END", "ANY",
            "HOW", "WHO", "WHY", "DID", "GET", "GOT", "LET", "PUT",
            "RUN", "SAY", "SEE", "TRY", "WAY", "DAY",
        }

        # Gather all acronyms from content
        found_acronyms = set()
        for s in sections:
            content = s.get("content") or ""
            matches = re.findall(r"\b([A-Z]{2,6})\b", content)
            found_acronyms.update(matches)

        found_acronyms -= non_acronyms

        # Load existing acronyms from DB
        existing_rows = conn.execute(
            "SELECT acronym, expansion FROM acronyms"
        ).fetchall()
        existing_map = {r["acronym"]: r["expansion"] for r in existing_rows}

        acronym_list = []
        new_count = 0

        for acr in sorted(found_acronyms):
            if acr in existing_map:
                acronym_list.append({
                    "acronym": acr,
                    "expansion": existing_map[acr],
                    "status": "known",
                })
            else:
                # Store as unknown — human must provide expansion
                acr_id = _acr_id()
                conn.execute(
                    "INSERT OR IGNORE INTO acronyms "
                    "(id, acronym, expansion, domain, usage_count, "
                    " created_at) "
                    "VALUES (?, ?, ?, ?, 1, ?)",
                    (acr_id, acr, f"[EXPANSION NEEDED: {acr}]",
                     "proposal", _now()),
                )
                acronym_list.append({
                    "acronym": acr,
                    "expansion": f"[EXPANSION NEEDED: {acr}]",
                    "status": "new",
                })
                new_count += 1

        _audit(conn, "production.acronyms",
               f"Extracted {len(acronym_list)} acronyms from {proposal_id}",
               "proposal", proposal_id,
               {"total": len(acronym_list), "new": new_count})
        conn.commit()

        return {
            "proposal_id": proposal_id,
            "acronyms": acronym_list,
            "total_count": len(acronym_list),
            "known_count": len(acronym_list) - new_count,
            "new_count": new_count,
        }
    finally:
        conn.close()


def check_cross_references(proposal_id, db_path=None):
    """Find broken or potentially invalid cross-references in proposal.

    Scans section content for patterns like "Section X.X", "Table X",
    "Figure X" and validates that referenced section numbers exist in
    the proposal's section list.

    Args:
        proposal_id: Proposal ID to check.
        db_path: Optional database path override.

    Returns:
        dict with total_references, valid_references, broken_references.

    Raises:
        ValueError: If proposal not found.
    """
    conn = _get_db(db_path)
    try:
        proposal = _load_proposal(conn, proposal_id)
        sections = _load_sections(conn, proposal_id)

        # Build set of known section numbers
        known_sections = {s["section_number"] for s in sections}

        all_refs = []
        broken_refs = []
        valid_refs = []

        for s in sections:
            content = s.get("content") or ""

            # Find "Section X.X.X" references
            sec_refs = re.findall(
                r"[Ss]ection\s+(\d+(?:\.\d+)*)", content
            )
            for ref in sec_refs:
                ref_entry = {
                    "source_section": s["section_number"],
                    "reference_type": "section",
                    "reference_target": ref,
                }
                all_refs.append(ref_entry)
                if ref in known_sections:
                    valid_refs.append(ref_entry)
                else:
                    ref_entry["issue"] = (
                        f"Section {ref} referenced but not found"
                    )
                    broken_refs.append(ref_entry)

            # Find "Table X" and "Figure X" references (informational)
            table_refs = re.findall(
                r"[Tt]able\s+(\d+(?:\.\d+)?(?:-\d+)?)", content
            )
            for ref in table_refs:
                all_refs.append({
                    "source_section": s["section_number"],
                    "reference_type": "table",
                    "reference_target": f"Table {ref}",
                })

            fig_refs = re.findall(
                r"[Ff]igure\s+(\d+(?:\.\d+)?(?:-\d+)?)", content
            )
            for ref in fig_refs:
                all_refs.append({
                    "source_section": s["section_number"],
                    "reference_type": "figure",
                    "reference_target": f"Figure {ref}",
                })

        return {
            "proposal_id": proposal_id,
            "total_references": len(all_refs),
            "valid_references": len(valid_refs),
            "broken_references": broken_refs,
            "broken_count": len(broken_refs),
            "all_references": all_refs,
        }
    finally:
        conn.close()


def get_package_status(proposal_id, db_path=None):
    """Get current packaging status for a proposal.

    Aggregates section readiness, review status, and CAG clearance
    into a single status view.

    Args:
        proposal_id: Proposal ID.
        db_path: Optional database path override.

    Returns:
        dict with section_readiness, review_status, cag_status,
        and overall packaging_ready flag.

    Raises:
        ValueError: If proposal not found.
    """
    conn = _get_db(db_path)
    try:
        proposal = _load_proposal(conn, proposal_id)
        sections = _load_sections(conn, proposal_id)

        # Section readiness
        status_counts = {}
        for s in sections:
            st = s.get("status") or "outline"
            status_counts[st] = status_counts.get(st, 0) + 1
        all_final = status_counts.get("final", 0) == len(sections)

        # Review status
        reviews = {}
        for rt in ("pink", "red", "gold", "white"):
            row = conn.execute(
                "SELECT overall_score, review_status, reviewed_at "
                "FROM proposal_reviews "
                "WHERE proposal_id = ? AND review_type = ? "
                "ORDER BY reviewed_at DESC LIMIT 1",
                (proposal_id, rt),
            ).fetchone()
            if row:
                reviews[rt] = {
                    "score": row["overall_score"],
                    "status": row["review_status"],
                    "reviewed_at": row["reviewed_at"],
                }
            else:
                reviews[rt] = {"score": None, "status": "not_started"}

        # CAG status
        cag = proposal.get("cag_status") or "pending"
        open_alerts = conn.execute(
            "SELECT COUNT(*) as cnt FROM cag_alerts "
            "WHERE proposal_id = ? AND status IN ('open', 'quarantined')",
            (proposal_id,),
        ).fetchone()
        alert_count = open_alerts["cnt"] if open_alerts else 0

        # Overall readiness
        reviews_complete = all(
            v["status"] == "completed" for v in reviews.values()
        )
        cag_clear = cag in ("clear",) and alert_count == 0
        packaging_ready = all_final and reviews_complete and cag_clear

        return {
            "proposal_id": proposal_id,
            "proposal_status": proposal.get("status"),
            "section_readiness": {
                "total_sections": len(sections),
                "status_breakdown": status_counts,
                "all_final": all_final,
            },
            "review_status": reviews,
            "cag_status": {
                "status": cag,
                "open_alerts": alert_count,
                "clear": cag_clear,
            },
            "packaging_ready": packaging_ready,
        }
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# Word (.docx) Export
# ---------------------------------------------------------------------------

def _configure_docx_styles(doc, classification):
    """Configure Word document styles for DoD proposal formatting.

    Sets up Heading 1/2/3 styles, Normal text, and header/footer
    with classification banners.
    """
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Times New Roman"
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    # Set margins (1 inch all sides — standard DoD)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Header with classification banner
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = classification
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.style.font.size = Pt(10)
        hp.style.font.bold = True
        hp.style.font.color.rgb = RGBColor(128, 0, 0)

        # Footer with classification + page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = classification
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.size = Pt(10)
        fp.style.font.bold = True
        fp.style.font.color.rgb = RGBColor(128, 0, 0)


def _write_volume_docx(proposal, vol_name, vol_sections, output_path,
                        classification, sol_num, template_path=None):
    """Write a single volume as a formatted Word (.docx) document.

    Args:
        proposal: Proposal dict.
        vol_name: Volume name (technical, management, etc.).
        vol_sections: List of section dicts for this volume.
        output_path: Path to write the .docx file.
        classification: Classification marking string.
        sol_num: Solicitation number.
        template_path: Optional custom .docx template path.
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required for Word export. "
            "Install with: pip install python-docx"
        )

    # Create document from template or blank
    if template_path and Path(template_path).exists():
        doc = python_docx.Document(str(template_path))
    else:
        doc = python_docx.Document()

    _configure_docx_styles(doc, classification)

    # --- Title Page ---
    doc.add_paragraph("")  # Spacer
    doc.add_paragraph("")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(proposal.get("title", "PROPOSAL"))
    run.font.size = Pt(24)
    run.font.bold = True

    doc.add_paragraph("")
    vol_para = doc.add_paragraph()
    vol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = vol_para.add_run(f"Volume: {vol_name.upper().replace('_', ' ')}")
    run.font.size = Pt(18)
    run.font.bold = True

    if sol_num:
        doc.add_paragraph("")
        sol_para = doc.add_paragraph()
        sol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sol_para.add_run(f"Solicitation: {sol_num}")
        run.font.size = Pt(14)

    doc.add_paragraph("")
    class_para = doc.add_paragraph()
    class_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = class_para.add_run(classification)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(128, 0, 0)

    # Page break after title
    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    for s in vol_sections:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_after = Pt(2)
        run = toc_para.add_run(
            f"{s['section_number']}    {s['section_title']}"
        )
        run.font.size = Pt(11)

    doc.add_page_break()

    # --- Section Content ---
    for s in vol_sections:
        doc.add_heading(
            f"{s['section_number']} {s['section_title']}",
            level=1,
        )

        content = s.get("content") or "[CONTENT MISSING]"

        # Split content into paragraphs and add them
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                doc.add_paragraph("")
                continue

            # Detect sub-headings (lines that are all caps or end with colon)
            if (para_text.isupper() and len(para_text) < 100) or \
               (para_text.endswith(":") and len(para_text) < 80):
                doc.add_heading(para_text, level=2)
            else:
                doc.add_paragraph(para_text)

    # Save
    doc.save(str(output_path))


def _write_acronym_table_docx(acronym_list, output_path, classification):
    """Write acronym list as a formatted Word table.

    Args:
        acronym_list: List of dicts with 'acronym' and 'expansion'.
        output_path: Path to write the .docx file.
        classification: Classification marking string.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx required for Word export")

    doc = python_docx.Document()
    _configure_docx_styles(doc, classification)

    doc.add_heading("Acronym List", level=1)

    if acronym_list:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Acronym"
        hdr_cells[1].text = "Expansion"
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for acr in sorted(acronym_list, key=lambda x: x.get("acronym", "")):
            row_cells = table.add_row().cells
            row_cells[0].text = acr.get("acronym", "")
            row_cells[1].text = acr.get("expansion", "")

    doc.save(str(output_path))


def _write_compliance_matrix_docx(matrix_entries, output_path, classification):
    """Write compliance matrix as a formatted Word table.

    Args:
        matrix_entries: List of compliance matrix entry dicts.
        output_path: Path to write the .docx file.
        classification: Classification marking string.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx required for Word export")

    doc = python_docx.Document()
    _configure_docx_styles(doc, classification)

    doc.add_heading("Compliance Matrix", level=1)

    if matrix_entries:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header
        headers = ["Req ID", "Source", "Requirement", "Status", "Volume/Section"]
        for i, hdr in enumerate(headers):
            table.rows[0].cells[i].text = hdr
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for entry in matrix_entries:
            row = table.add_row().cells
            row[0].text = entry.get("requirement_id", "")
            row[1].text = entry.get("source", "")
            row[2].text = (entry.get("requirement_text", ""))[:150]
            row[3].text = entry.get("compliance_status", "not_addressed")
            sec_info = entry.get("section_number") or ""
            if entry.get("volume"):
                sec_info = f"{entry['volume']}/{sec_info}" if sec_info else entry["volume"]
            row[4].text = sec_info

            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    doc.save(str(output_path))


def package_proposal_docx(proposal_id, output_dir, template_path=None,
                           db_path=None):
    """Package proposal as formatted Word (.docx) documents.

    Produces one .docx file per volume, plus acronym list and compliance
    matrix as separate Word documents.

    Args:
        proposal_id: Proposal ID.
        output_dir: Directory to write .docx files.
        template_path: Optional custom .docx template.
        db_path: Override database path.

    Returns:
        dict with file manifest and packaging status.
    """
    if not HAS_DOCX:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    conn = _get_db(db_path)
    try:
        proposal = _load_proposal(conn, proposal_id)
        sections = _load_sections(conn, proposal_id)
        if not sections:
            raise ValueError(f"No sections found for proposal {proposal_id}")

        classification = proposal.get("classification") or "CUI // SP-PROPIN"
        opp_id = proposal.get("opportunity_id") or "UNKNOWN"
        sol_num = ""
        if opp_id != "UNKNOWN":
            opp_row = conn.execute(
                "SELECT solicitation_number FROM opportunities WHERE id = ?",
                (opp_id,),
            ).fetchone()
            if opp_row:
                sol_num = opp_row["solicitation_number"] or ""

        out_path = Path(output_dir)
        out_path.mkdir(parents=True, exist_ok=True)

        sol_slug = _safe_filename(sol_num) if sol_num else "proposal"
        files_written = []

        # Group sections by volume
        volumes_defined = _parse_json_field(proposal.get("volumes")) or [
            "technical", "management", "past_performance", "cost"
        ]
        volume_sections = {}
        for s in sections:
            vol = s.get("volume", "technical")
            volume_sections.setdefault(vol, []).append(s)

        # Write each volume as a .docx
        for vol_name in volumes_defined:
            vol_secs = volume_sections.get(vol_name, [])
            if not vol_secs:
                continue

            filename = f"{sol_slug}_vol_{vol_name}.docx"
            filepath = out_path / filename

            _write_volume_docx(
                proposal, vol_name, vol_secs, filepath,
                classification, sol_num, template_path=template_path,
            )

            files_written.append({
                "filename": filename,
                "volume": vol_name,
                "sections": len(vol_secs),
                "size_bytes": filepath.stat().st_size,
                "format": "docx",
                "path": str(filepath),
            })

        # Write acronym list as Word table
        acronyms = generate_acronym_list(proposal_id, db_path=db_path)
        if acronyms.get("acronyms"):
            acr_filename = f"{sol_slug}_acronyms.docx"
            acr_path = out_path / acr_filename
            _write_acronym_table_docx(
                acronyms["acronyms"], acr_path, classification,
            )
            files_written.append({
                "filename": acr_filename,
                "volume": "attachments",
                "sections": 0,
                "size_bytes": acr_path.stat().st_size,
                "format": "docx",
                "path": str(acr_path),
            })

        # Write compliance matrix as Word table
        matrix_rows = conn.execute(
            "SELECT * FROM compliance_matrices WHERE proposal_id = ? "
            "ORDER BY source, requirement_id",
            (proposal_id,),
        ).fetchall()
        if matrix_rows:
            matrix_entries = [dict(r) for r in matrix_rows]
            mx_filename = f"{sol_slug}_compliance_matrix.docx"
            mx_path = out_path / mx_filename
            _write_compliance_matrix_docx(
                matrix_entries, mx_path, classification,
            )
            files_written.append({
                "filename": mx_filename,
                "volume": "attachments",
                "sections": 0,
                "size_bytes": mx_path.stat().st_size,
                "format": "docx",
                "path": str(mx_path),
            })

        result = {
            "proposal_id": proposal_id,
            "package_id": _pkg_id(),
            "output_format": "docx",
            "output_dir": str(out_path),
            "files": files_written,
            "total_files": len(files_written),
            "total_size_bytes": sum(f["size_bytes"] for f in files_written),
            "acronym_count": len(acronyms.get("acronyms", [])),
            "compliance_matrix_entries": len(matrix_rows) if matrix_rows else 0,
            "packaged_at": _now(),
        }

        _audit(conn, "production.package_docx",
               f"Packaged proposal {proposal_id} as DOCX",
               "proposal", proposal_id,
               {"files": len(files_written), "format": "docx"})
        conn.commit()
        return result
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _build_parser():
    """Build argument parser for the CLI."""
    import argparse
    parser = argparse.ArgumentParser(
        description="GovProposal Submission Packager",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  %(prog)s --package --proposal-id PROP-001 "
            "--output /tmp/out --json\n"
            "  %(prog)s --package --proposal-id PROP-001 "
            "--output /tmp/out --format docx --json\n"
            "  %(prog)s --package --proposal-id PROP-001 "
            "--output /tmp/out --format docx --template custom.docx\n"
            "  %(prog)s --validate --proposal-id PROP-001 --json\n"
            "  %(prog)s --acronyms --proposal-id PROP-001 --json\n"
            "  %(prog)s --cross-refs --proposal-id PROP-001 --json\n"
            "  %(prog)s --status --proposal-id PROP-001 --json\n"
        ),
    )

    action = parser.add_mutually_exclusive_group(required=True)
    action.add_argument("--package", action="store_true",
                        help="Assemble final proposal package")
    action.add_argument("--validate", action="store_true",
                        help="Pre-submission validation only")
    action.add_argument("--acronyms", action="store_true",
                        help="Extract and list acronyms")
    action.add_argument("--cross-refs", action="store_true",
                        help="Check cross-references")
    action.add_argument("--status", action="store_true",
                        help="Get packaging status")

    parser.add_argument("--proposal-id", help="Proposal ID")
    parser.add_argument("--output", help="Output directory (for --package)")
    parser.add_argument("--format", choices=["txt", "docx"], default="txt",
                        help="Output format (default: txt)")
    parser.add_argument("--template",
                        help="Custom .docx template path (for --format docx)")
    parser.add_argument("--json", action="store_true", help="JSON output")
    parser.add_argument("--db-path", help="Override database path")

    return parser


def main():
    """CLI entry point."""
    parser = _build_parser()
    args = parser.parse_args()
    db = args.db_path

    try:
        if not args.proposal_id:
            parser.error("--proposal-id is required")

        if args.package:
            if not args.output:
                parser.error("--package requires --output")
            if args.format == "docx":
                result = package_proposal_docx(
                    args.proposal_id, args.output,
                    template_path=args.template, db_path=db,
                )
            else:
                result = package_proposal(
                    args.proposal_id, args.output, db_path=db
                )

        elif args.validate:
            result = validate_submission(args.proposal_id, db_path=db)

        elif args.acronyms:
            result = generate_acronym_list(args.proposal_id, db_path=db)

        elif args.cross_refs:
            result = check_cross_references(args.proposal_id, db_path=db)

        elif args.status:
            result = get_package_status(args.proposal_id, db_path=db)

        # Output
        if args.json:
            print(json.dumps(result, indent=2, default=str))
        else:
            if isinstance(result, dict):
                for key, value in result.items():
                    if isinstance(value, (list, dict)):
                        if key == "checklist":
                            print(f"  {key}:")
                            for item in value:
                                mark = "[X]" if item["passed"] else "[ ]"
                                print(f"    {mark} {item['item']}")
                        elif key == "files":
                            print(f"  {key}: ({len(value)} files)")
                            for f in value:
                                print(f"    {f['filename']} "
                                      f"({f['size_bytes']} bytes)")
                        elif key == "acronyms":
                            print(f"  {key}: ({len(value)} entries)")
                        else:
                            print(f"  {key}: "
                                  f"{json.dumps(value, default=str)}")
                    else:
                        print(f"  {key}: {value}")

    except ValueError as exc:
        if args.json:
            print(json.dumps({"error": str(exc)}, indent=2))
        else:
            print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)
    except sqlite3.Error as exc:
        if args.json:
            print(json.dumps({"error": f"Database error: {exc}"}, indent=2))
        else:
            print(f"Database error: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
